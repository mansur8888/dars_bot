import asyncio
import os
import docx
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup, InlineKeyboardButton, FSInputFile
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext

TOKEN = "8685388983:AAFwjfV-RvOrq4vT1hI_SxIqIPd0-lZ6cZg"

bot = Bot(token=TOKEN)
dp = Dispatcher()

class DarsTaqsimot(StatesGroup):
    subject = State()
    total_hours = State()
    teacher_count = State()
    teacher_name = State()
    teacher_step = State()
    teacher_category = State()
    teacher_has_cert = State()
    teacher_cert_subject = State()
    teacher_cert_type = State()
    teacher_cert_level = State()
    teacher_retrain = State()
    teacher_side = State()
    teacher_student = State()
    teacher_middle = State()
    teacher_leader = State()

SUBJECTS_YUQORI = [
    "Ona tili", "Adabiyot", "Rus tili", "Chet tili", "Tarix", 
    "Davlat va huquq asoslari", "Tarbiya", "Matematika", 
    "Informatika va axborot texnologiyalari", "Fizika", "Astronomiya", 
    "Kimyo", "Biologiya", "Geografiya", "Iqtisodiy bilim asoslari", 
    "Tadbirkorlik asoslari", "Tabiiy fan (Science)", "Musiqa madaniyati", 
    "Tasviriy san'at", "Chizmachilik", "Texnologiya", "Jismoniy tarbiya", 
    "Chaqiruvga qadar boshlang'ich tayyorgarlik"
]

main_kb = ReplyKeyboardMarkup(
    keyboard=[
        [KeyboardButton(text="🚀 Botni boshlash / Restart")],
        [KeyboardButton(text="📚 Dars soatini taqsimlash boshlash")],
        [KeyboardButton(text="📜 Nizom bo'yicha ketma-ketlik")]
    ],
    resize_keyboard=True
)

@dp.message(Command("start"))
@dp.message(F.text.in_(["🚀 Botni boshlash / Restart", "📚 Dars soatini taqsimlash boshlash"]))
async def start_cmd(message: types.Message, state: FSMContext):
    await state.clear()
    kb_list = [[InlineKeyboardButton(text=sub, callback_data=f"sub_{i}")] for i, sub in enumerate(SUBJECTS_YUQORI)]
    kb_list.append([InlineKeyboardButton(text="✅ Tayyor", callback_data="sub_ready")])
    kb = InlineKeyboardMarkup(inline_keyboard=kb_list)
    
    await state.set_state(DarsTaqsimot.subject)
    await message.answer(
        "<b>3271-sonli Qaror va Nizom asosida Dars soatlarini taqsimlash botiga xush kelibsiz!</b>\n\n"
        "Iltimos, fanni tanlang:",
        parse_mode="HTML",
        reply_markup=kb
    )

@dp.message(F.text == "📜 Nizom bo'yicha ketma-ketlik")
async def info_nizom(message: types.Message):
    text = (
        "<b>Dars soatlarini taqsimlashning rasmiy ketma-ketligi:</b>\n\n"
        "1️⃣ Oliy malaka toifasiga ega o'qituvchilar;\n"
        "2️⃣ Sertifikatga ega o'qituvchilar;\n"
        "3️⃣ Birinchi malaka toifasiga ega o'qituvchilar;\n"
        "4️⃣ Ikkinchi malaka toifasiga ega o'qituvchilar;\n"
        "5️⃣ Toifasiz o'qituvchilar;\n"
        "6️⃣ Qayta tayyorlash diplomiga ega o'qituvchilar;\n"
        "7️⃣ O'rindoshlar;\n"
        "8️⃣ Talabalar;\n"
        "9️⃣ O'rta maxsus / professional ma'lumotlilar;\n"
        "🔟 Rahbar xodimlar."
    )
    await message.answer(text, parse_mode="HTML", reply_markup=main_kb)

@dp.callback_query(DarsTaqsimot.subject, F.data.startswith("sub_"))
async def select_subject(callback: types.CallbackQuery, state: FSMContext):
    if callback.data == "sub_ready":
        await callback.message.answer("Iltimos, ro'yxatdan fanni tanlang.")
        await callback.answer()
        return
    
    idx = int(callback.data.split("_")[1])
    selected_sub = SUBJECTS_YUQORI[idx]
    
    await state.update_data(subject_name=selected_sub)
    await state.set_state(DarsTaqsimot.total_hours)
    
    await callback.message.edit_text(
        f"<b>Fan:</b> {selected_sub}\n\n"
        f"Ushbu fan bo'yicha maktabdagi <b>jami dars soatini</b> kiriting (Faqat raqam):",
        parse_mode="HTML"
    )
    await callback.answer()

@dp.message(DarsTaqsimot.total_hours)
async def process_total_hours(message: types.Message, state: FSMContext):
    if not message.text.isdigit():
        await message.answer("Iltimos, faqat raqam kiriting:", reply_markup=main_kb)
        return
    
    total = int(message.text)
    await state.update_data(total_hours=total, teachers=[], current_teacher=1)
    await state.set_state(DarsTaqsimot.teacher_name)
    
    await message.answer("<b>1-o'qituvchi ma'lumotlarini kiriting.</b>\nIsmingizni kiriting:", parse_mode="HTML", reply_markup=types.ReplyKeyboardRemove())

@dp.message(DarsTaqsimot.teacher_name)
async def process_teacher_name(message: types.Message, state: FSMContext):
    t_name = message.text.strip()
    await state.update_data(current_teacher_name=t_name)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Oliy ma'lumot", callback_data="edu_oliy")],
        [InlineKeyboardButton(text="O'rta maxsus/professional", callback_data="edu_orta")],
        [InlineKeyboardButton(text="Talaba", callback_data="edu_talaba")]
    ])
    await state.set_state(DarsTaqsimot.teacher_category)
    await message.answer("Ta'lim darajasi", reply_markup=kb)

@dp.callback_query(DarsTaqsimot.teacher_category, F.data.startswith("edu_"))
async def process_edu_level(callback: types.CallbackQuery, state: FSMContext):
    edu = callback.data.replace("edu_", "")
    await state.update_data(edu_level=edu)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Boshlang'ich sinf o'qituvchisi (1-4-sinf)", callback_data="cls_bosh")],
        [InlineKeyboardButton(text="Yuqori sinf (fan) o'qituvchisi (5-11-sinf)", callback_data="cls_yuqori")]
    ])
    await callback.message.edit_text("Siz boshlang'ich sinf o'qituvchimisiz yoki yuqori sinf (fan) o'qituvchimisiz?", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.in_(["cls_bosh", "cls_yuqori"]))
async def process_class_type(callback: types.CallbackQuery, state: FSMContext):
    # Agar boshlang'ich sinf tanlansa 1 stavka = 18 soat, yuqori sinf uchun = 20 soat
    is_bosh = (callback.data == "cls_bosh")
    rate = 18 if is_bosh else 20
    await state.update_data(rate=rate)

    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Oliy toifa", callback_data="cat_oliy")],
        [InlineKeyboardButton(text="1-toifa", callback_data="cat_1")],
        [InlineKeyboardButton(text="2-toifa", callback_data="cat_2")],
        [InlineKeyboardButton(text="Toifasiz", callback_data="cat_toifasiz")]
    ])
    await callback.message.edit_text("Malaka toifangiz", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.startswith("cat_"))
async def process_category(callback: types.CallbackQuery, state: FSMContext):
    cat = callback.data.replace("cat_", "")
    await state.update_data(category_tier=cat)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Ha", callback_data="cert_ha")],
        [InlineKeyboardButton(text="Yo'q", callback_data="cert_yoq")]
    ])
    await callback.message.edit_text("Dars beradigan faningiz bo'yicha milliy yoki xalqaro tan olingan sertifikatingiz bormi?", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.startswith("cert_"))
async def process_cert(callback: types.CallbackQuery, state: FSMContext):
    cert = callback.data.replace("cert_", "")
    await state.update_data(has_cert=cert)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Ha", callback_data="ret_ha")],
        [InlineKeyboardButton(text="Yo'q", callback_data="ret_yoq")]
    ])
    await callback.message.edit_text("Qayta tayyorlash bo'yicha diplomingiz bormi?", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.startswith("ret_"))
async def process_retrain(callback: types.CallbackQuery, state: FSMContext):
    ret = callback.data.replace("ret_", "")
    await state.update_data(has_retrain=ret)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Ha", callback_data="side_ha")],
        [InlineKeyboardButton(text="Yo'q", callback_data="side_yoq")]
    ])
    await callback.message.edit_text("O'rindoshlik asosida ishlaysizmi?", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.startswith("side_"))
async def process_side(callback: types.CallbackQuery, state: FSMContext):
    side = callback.data.replace("side_", "")
    await state.update_data(is_side=side)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Ha", callback_data="stud_ha")],
        [InlineKeyboardButton(text="Yo'q", callback_data="stud_yoq")]
    ])
    await callback.message.edit_text("Talabamisiz?", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.startswith("stud_"))
async def process_student(callback: types.CallbackQuery, state: FSMContext):
    stud = callback.data.replace("stud_", "")
    await state.update_data(is_student=stud)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Ha", callback_data="lead_ha")],
        [InlineKeyboardButton(text="Yo'q", callback_data="lead_yoq")]
    ])
    await callback.message.edit_text("Rahbar xodimmmisiz?", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.startswith("lead_"))
async def process_leader(callback: types.CallbackQuery, state: FSMContext):
    lead = callback.data.replace("lead_", "")
    data = await state.get_data()
    
    cat = data.get('category_tier')
    cert = data.get('has_cert')
    ret = data.get('has_retrain')
    side = data.get('is_side')
    stud = data.get('is_student')
    
    step = 10
    status_name = "Rahbar xodim"
    
    if lead == "yoq":
        if cat == "oliy":
            step = 1
            status_name = "Oliy toifali o'qituvchi"
        elif cert == "ha":
            step = 2
            status_name = "Sertifikatli o'qituvchi"
        elif cat == "1":
            step = 3
            status_name = "1-toifali o'qituvchi"
        elif cat == "2":
            step = 4
            status_name = "2-toifali o'qituvchi"
        elif cat == "toifasiz" and side == "yoq" and stud == "yoq" and ret == "yoq":
            step = 5
            status_name = "Toifasiz mutaxassis"
        elif ret == "ha":
            step = 6
            status_name = "Qayta tayyorlash diplomiga ega"
        elif side == "ha":
            step = 7
            status_name = "O'rindosh o'qituvchi"
        elif stud == "ha":
            step = 8
            status_name = "Talaba o'qituvchi"
        else:
            step = 9
            status_name = "O'rta maxsus o'qituvchi"

    teachers = data.get('teachers', [])
    teachers.append({
        "id": data['current_teacher'],
        "fish": data['current_teacher_name'],
        "step": step,
        "status_name": status_name
    })
    
    current = data['current_teacher']
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Yana o'qituvchi qo'shish", callback_data="add_more_yes")],
        [InlineKeyboardButton(text="O'qituvchilar tugadi", callback_data="add_more_no")]
    ])
    
    await state.update_data(teachers=teachers)
    await callback.message.edit_text(f"Qabul qilindi. Hozircha {current} ta o'qituvchi kiritildi.", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data == "add_more_yes")
async def add_more_teacher(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    next_num = data['current_teacher'] + 1
    await state.update_data(current_teacher=next_num)
    await state.set_state(DarsTaqsimot.teacher_name)
    await callback.message.answer(f"<b>{next_num}-o'qituvchining ismini kiriting:</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "add_more_no")
async def finish_teachers(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    teachers = data['teachers']
    
    rate = data.get('rate', 20) # Boshlang'ich uchun 18, yuqori uchun 20
    max_limit = int(rate * 1.5) # 1.5 stavka limit (masalan, 20 soatlik uchun 30 soat)
    total_hours = int(data['total_hours'])
    
    # Nizom ketma-ketligi bo'yicha saralash
    sorted_teachers = sorted(teachers, key=lambda x: x['step'])
    
    remaining_hours = total_hours
    teacher_assignments = {t['id']: 0 for t in sorted_teachers}

    # 1-bosqich: Har bir o'qituvchiga NAVBAT BILAN o'z limitigacha (1.5 stavkagacha) to'liq soat berib borish
    for t in sorted_teachers:
        if remaining_hours <= 0:
            break
        can_take = max_limit - teacher_assignments[t['id']]
        give = min(remaining_hours, can_take)
        teacher_assignments[t['id']] += give
        remaining_hours -= give

    # 2-bosqich: Agar shundan keyin ham dars ortib qolsa (masalan, o'qituvchilar kam bo'lib, jami soat juda ko'p bo'lsa), 
    # cheklovni olib tashlab yoki taqsimotni davom ettirish mumkin, lekin hozirgi holatda 1.5 stavka limit yetarli bo'ladi.

    results = [(t, teacher_assignments[t['id']]) for t in sorted_teachers]

    res_text = (
        f"📊 <b>DARS TAQSIMOTI NATIJASI</b>\n"
        f"(RASMIY NIZOM BO'YICHA)\n\n"
        f"🔹 <b>Fan:</b> {data['subject_name']}\n"
        f"🔹 <b>Jami dars soati:</b> {total_hours} soat\n"
        f"🔹 <b>1 stavka me'yori:</b> {rate} soat\n"
        f"🔹 <b>O'qituvchilar soni:</b> {len(teachers)} nafar\n\n"
        f"<b>Ustuvor ketma-ketlik bo'yicha taqsimot:</b>\n"
    )
    
    for idx, (t, h) in enumerate(results, 1):
        stavka = round(h / rate, 2)
        res_text += (
            f"\n<b>{idx}. {t['fish']}</b>\n"
            f" └ Maqomi: <i>{t['status_name']} ({t['step']}-o'rin)</i>\n"
            f" └ Dars soati: <b>{h} soat</b> ({stavka} stavka)\n"
        )
    
    if remaining_hours == 0:
        res_text += f"\n✅ <b>Barcha dars soatlari to'liq taqsimlandi!</b>"
    else:
        res_text += f"\n⚠️ <b>Taqsimlanmay qolgan dars soati:</b> {remaining_hours} soat."

    doc = docx.Document()
    doc.add_heading('DARS SOATLARINI TAQSIMLASH BAYONNOMASI', 0)
    doc.add_paragraph(f"Fan: {data['subject_name']}")
    doc.add_paragraph(f"Jami dars soati: {total_hours} soat")
    
    for idx, (t, h) in enumerate(results, 1):
        stavka = round(h / rate, 2)
        doc.add_paragraph(f"{idx}. {t['fish']} - {t['status_name']}: {h} soat ({stavka} stavka)", style='List Bullet')
    
    file_path = f"dars_taqsimoti_{callback.message.chat.id}.docx"
    doc.save(file_path)

    await callback.message.edit_text(res_text, parse_mode="HTML")
    
    doc_file = FSInputFile(file_path)
    await callback.message.answer_document(doc_file, caption="📄 Tayyor rasmiy hujjat (Word formatida)", reply_markup=main_kb)
    
    if os.path.exists(file_path):
        os.remove(file_path)
        
    await state.clear()
    await callback.answer()

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
